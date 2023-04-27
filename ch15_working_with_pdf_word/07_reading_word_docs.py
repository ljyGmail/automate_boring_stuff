import docx

doc = docx.Document('data/ch15/demo.docx')

print(f'len(doc.paragraphs): {len(doc.paragraphs)}')

print(f'doc.paragraphs[0].text: {doc.paragraphs[0].text}')
print(f'doc.paragraphs[1].text: {doc.paragraphs[1].text}')

print(f'len(doc.paragraphs[1].runs): {len(doc.paragraphs[1].runs)}')

print(f'doc.paragraphs[1].runs[0].text: {doc.paragraphs[1].runs[0].text}')
print(f'doc.paragraphs[1].runs[1].text: {doc.paragraphs[1].runs[1].text}')
print(f'doc.paragraphs[1].runs[2].text: {doc.paragraphs[1].runs[2].text}')
print(f'doc.paragraphs[1].runs[3].text: {doc.paragraphs[1].runs[3].text}')
print(f'doc.paragraphs[1].runs[4].text: {doc.paragraphs[1].runs[4].text}')
