import docx

doc = docx.Document()

doc.add_picture('data/ch15/zophie.png',
                width=docx.shared.Inches(1), height=docx.shared.Cm(4))

doc.save('data/ch15/add_picture.docx')
