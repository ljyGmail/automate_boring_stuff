import docx

doc = docx.Document()

print(
    f"doc.add_paragraph('Hello, world!'): {doc.add_paragraph('Hello, world!')}")

doc.save('data/ch15/helloworld.docx')

doc = docx.Document()

doc.add_paragraph('Hello world!')

paraObj1 = doc.add_paragraph('This is a second paragraph.')
# pass an optional second argument that is a string of the Paragraph or Run object's style
paraObj2 = doc.add_paragraph('This is a yet another paragraph.', 'Title')

paraObj1.add_run(' This text is being added to the second paragraph.')

doc.save('data/ch15/multipleParagraphs.docx')
