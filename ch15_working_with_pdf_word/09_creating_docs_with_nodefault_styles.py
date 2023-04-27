import docx

doc = docx.Document('data/ch15/demo.docx')

print(f'doc.paragraphs[0].text: {doc.paragraphs[0].text}')

# The exact id may be different
print(f'doc.paragraphs[0].style: {doc.paragraphs[0].style}')

doc.paragraphs[0].style = 'Normal'

print(f'doc.paragraphs[1].text: {doc.paragraphs[1].text}')

print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text,
      doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text, doc.paragraphs[1].runs[4].text))

doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True

doc.save('data/ch15/restyled.docx')
