import docx

doc = docx.Document()

print(
    f"doc.add_paragraph('This is on the first page!'): {doc.add_paragraph('This is on the first page!')}")

doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.paragraphs[1].runs[0].add_break()
doc.paragraphs[1].add_run('This is one the second line')

doc.save('data/ch15/twoPage.docx')
